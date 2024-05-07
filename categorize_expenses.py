import pandas as pd
import os
import matplotlib.pyplot as plt
import docx
from docx import Document
from io import BytesIO
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

def categorize_expenditure(excel_file, description_column, amount_column,income_column):
    # Load Excel file into a DataFrame
    df = pd.read_excel(excel_file)

    # Replace NaN values in description and amount columns
    df[description_column] = df[description_column].fillna('')
    df[amount_column] = df[amount_column].fillna(0)
    df[income_column] = df[income_column].fillna(0)
    
    # Dictionary to store category-wise expenditure totals
    categories = {'Groceries': 0, 'Bills': 0, 'Entertainment': 0, 'Rent': 0,'Credit' : 0,  'Food': 0,'Salon': 0,'EMI': 0,'Others': 0,'Savings': 0,'Giving':0}
    totals = {'Total Income':0,'Total Saving':0,'Total Giving':0, 'Total Expenditure':0}

    # Categorize expenditures
    for index, row in df.iterrows():
        description = str(row[description_column]).lower()
        amount = row[amount_column]
        if 'groceries' in str(row[description_column]).lower():
            categories['Groceries'] += row[amount_column]
        elif 'bill' in str(row[description_column]).lower():
            categories['Bills'] += row[amount_column]
        elif 'entertainment' in str(row[description_column]).lower():
            categories['Entertainment'] += row[amount_column]
        elif 'rent' in str(row[description_column]).lower():
            categories['Rent'] += row[amount_column]
        elif 'credit' in str(row[description_column]).lower():
            categories['Credit'] += row[amount_column]
        elif 'food' in str(row[description_column]).lower():
            categories['Food'] += row[amount_column]
        elif 'salon' in str(row[description_column]).lower():
            categories['Salon'] += row[amount_column]
        elif 'emi' in str(row[description_column]).lower():
            categories['EMI'] += row[amount_column]
        elif 'saving' in str(row[description_column]).lower():
            categories['Savings'] += row[amount_column]
            totals['Total Saving'] += row[amount_column] 
        elif 'giving' in str(row[description_column]).lower():
            categories['Giving'] += row[amount_column]
            totals['Total Giving'] += row[amount_column]
        elif 'salary' in str(row[description_column]).lower():
            totals['Total Income'] += row[income_column]
        else:
            categories['Others'] += float(row[amount_column])

    # Calculate total expenditure
    
    total_expenditure = sum(categories.values()) - categories['Giving'] - categories['Savings']
    totals['Total Expenditure'] = total_expenditure 
    
    

    # Write category-wise expenditure totals to a text file
    with open('expenditure_summary.txt', 'w') as file:
        for category, total in categories.items():
            file.write(f"{category}: {total:.2f} Rupees\n")
        for totals, indtotals in totals.items():
            file.write(f"{totals} : {indtotals:.2f} Rupees\n")




def visualize_expenditure_summary(text_file, docx_file):
    # Read data from text file
    categories = []
    amounts = []
    with open(text_file, 'r') as file:
        for line in file:
            if line.startswith('Total Income'):
                total_income = float(re.search(r'\d+\.\d+', line).group())
            elif line.startswith('Total Saving'):
                total_savings = float(re.search(r'\d+\.\d+', line).group())
            elif line.startswith('Total Expenditure'):
                total_expenditure = float(re.search(r'\d+\.\d+', line).group())
            elif line.startswith('Total Giving'):
                total_giving = float(re.search(r'\d+\.\d+', line).group())
            else:
                category, amount = line.strip().split(': ')
                categories.append(category)
                amounts.append(float(amount.replace('Rupees', '')))

    # Sort categories and amounts based on amounts
    sorted_data = sorted(zip(categories, amounts), key=lambda x: x[1])

    # Extract sorted categories and amounts
    sorted_categories = [item[0] for item in sorted_data]
    sorted_amounts = [item[1] for item in sorted_data]

    # Plot bar chart
    plt.figure(figsize=(10, 6))
    colors = ['yellow' if category.lower() == 'emi' else 'green' if category.lower() == 'savings' else 'red' for category in sorted_categories]
    plt.barh(sorted_categories, sorted_amounts, color=colors)
    plt.xlabel('Expenditure')
    plt.ylabel('Categories')
    plt.title('Expenditure Summary')
    plt.gca().invert_yaxis()  # Invert y-axis to have the largest category at the top


    # Save the bar chart image to a file
    bar_chart_file = 'bar_chart.png'
    plt.savefig(bar_chart_file, format='png')
    plt.close()

    # Create a pie chart for total income, savings, and expenditure
    labels = ['Total Savings', 'Spendings','Giving','EMI/Loans/Rent']
    sizes = [ total_savings, total_expenditure - amounts[10]-amounts[7]-amounts[3],total_giving,(float(amounts[7])+float(amounts[3]))]
    colors = [ 'green', 'red','skyblue','yellow']
    
    plt.figure(figsize=(6, 6))
    plt.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
    plt.title('Summary')

    # Save the pie chart image to a file
    pie_chart_file = 'pie_chart.png'
    plt.savefig(pie_chart_file, format='png')
    plt.close()

    # Create a new Word document
    doc = Document()
    
    # Add text content to the document
    doc.add_heading('Expenditure Summary', level=1).alignment = 1

    # Add a table for expenditure data
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'  # Set table style to have grid lines
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Center align the table

    # Add bold formatting to column headings
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Amount'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)

    
    for category, amount in zip(categories, amounts):
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = '₹{:.2f}'.format(amount)  # Convert amount to string with currency format

    # Add total income, savings, and expenditure to the table
    #total_expenditure = sum(amounts)
    row_cells = table.add_row().cells
    row_cells[0].text = 'Total Giving'
    row_cells[1].text = '₹{:.2f}'.format(total_giving)
    row_cells[0].paragraphs[0].runs[0].font.bold = True  # Set font to bold
    row_cells[1].paragraphs[0].runs[0].font.bold = True  # Set font to bold
    row_cells[0].text = 'Total Income'
    row_cells[1].text = '₹{:.2f}'.format(total_income)
    row_cells[0].paragraphs[0].runs[0].font.bold = True  # Set font to bold
    row_cells[1].paragraphs[0].runs[0].font.bold = True  # Set font to bold
    row_cells = table.add_row().cells
    row_cells[0].text = 'Total Savings'
    row_cells[1].text = '₹{:.2f}'.format(total_savings)
    row_cells[0].paragraphs[0].runs[0].font.bold = True  # Set font to bold
    row_cells[1].paragraphs[0].runs[0].font.bold = True  # Set font to bold
    row_cells = table.add_row().cells
    row_cells[0].text = 'Total Expenditure'
    row_cells[1].text = '₹{:.2f}'.format(total_expenditure)
    row_cells[0].paragraphs[0].runs[0].font.bold = True  # Set font to bold
    row_cells[1].paragraphs[0].runs[0].font.bold = True  # Set font to bold

    # Add the bar chart image to the document
    doc.add_picture(bar_chart_file, width=Inches(6))  # Set the width of the image to 6 inches

    # Add the pie chart image to the document
    doc.add_picture(pie_chart_file, width=Inches(6))  # Set the width of the image to 6 inches

    # Save the Word document
    doc.save(docx_file)

# Example usage:


if __name__ == "__main__":
    # Get Excel file path from user input
    excel_file = input("Enter the path of the Excel file: ")

    # Check if the file exists
    if not os.path.exists(excel_file):
        print("File not found.")
    else:
        # Load Excel file into a DataFrame
        df = pd.read_excel(excel_file)
        
        # Print available column names
        # print("Available column names:")
        # print(df.columns.tolist())
        
        # Get column names from user input
        # description_column = input("Enter the name of the column containing descriptions: ")
        # amount_column = input("Enter the name of the column containing amounts: ")
        
        # Categorize expenditure and write to text file
        categorize_expenditure(excel_file, 'Remarks', 'Amount','Deposit Amt.')

        # Visualize expenditure summary
        visualize_expenditure_summary('expenditure_summary.txt','expenditure_summary.docx')

        print("Expenditure summary written to expenditure_summary.docx.")

        os.startfile('expenditure_summary.docx')
